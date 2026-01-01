# -*- coding: utf-8 -*-
"""
文件名称：utils/create_dummy_docx.py
主要作用：生成测试用 Word 文档
实现功能：
1. 使用 python-docx 创建一个包含标准结构的 Word 文档
2. 用于测试解析器和生成流程
3. 包含标题、列表、段落等多种格式
"""
from docx import Document

def create_dummy_docx(path):
    doc = Document()
    
    # --- Metadata (Preamble) ---
    doc.add_paragraph('公司名称：xx科技有限公司')
    doc.add_paragraph('项目名称：xxxx-软件项目')
    doc.add_paragraph('汇报人：zwh')
    doc.add_paragraph('部门 / 团队：研发部')
    doc.add_paragraph('日期：2025.12.18')
    
    # --- Section 1 ---
    doc.add_heading('一、项目背景与目标', level=1)
    doc.add_paragraph('当前项目进度在等待smt公司贴片上，smt公司贴片完成后需将开发板给到硬件调试，调试完成后才可以开始软件调试。当前软件已完成基础框架的搭建')
    
    doc.add_heading('项目目标', level=2)
    doc.add_paragraph('目标 1：开发板到手后两个星期完成基础外设adc、pwm等调试', style='List Bullet')
    doc.add_paragraph('目标 2：在一月底实现mppt算法模块的调试，确保mppt算法能达到“晟高-800w”优化器的80%水平', style='List Bullet')

    # --- Section 2 ---
    doc.add_heading('二、项目整体方案', level=1)
    doc.add_paragraph('目前是将优化器软件分为4个大的模块：mppt算法模块、wisun通讯模块、保护模块、OTA模块')
    
    doc.add_heading('系统架构 / 技术路线', level=2)
    doc.add_paragraph('总体架构说明：软件划分为4个模块主要是根据功能的独立性', style='List Bullet')
    doc.add_paragraph('核心模块划分：mppt算法模块、wisun通讯模块、保护模块、OTA模块', style='List Bullet')
    doc.add_paragraph('关键技术选型理由：当前采用mppt+电压环直接输出占空比，因为加上电流会整体会复杂几倍，先将简单的调试明白', style='List Bullet')

    doc.add_heading('工作范围', level=2)
    doc.add_paragraph('本人 / 本团队负责内容：本人负责四大模块的整体设计以及后续的测试方案的制定', style='List Bullet')
    doc.add_paragraph('与其他模块的接口关系：/', style='List Bullet')

    # --- Section 3 ---
    doc.add_heading('三、关键技术与实现难点', level=1)
    doc.add_paragraph('本项目的关键技术为mppt算法的选择和wisun通讯的协议栈')
    
    doc.add_heading('关键技术点', level=2)
    doc.add_paragraph('技术点 1：mppt算法初步考虑是采用简单可靠的变步长P&O-MPPT；', style='List Bullet')
    doc.add_paragraph('技术点 2：wisun通讯当前采用自定义协议，目前只有握手、数据帧（心跳）、故障帧、控制帧、应答帧', style='List Bullet')
    doc.add_paragraph('技术点 3：具备硬件过流保护和软件保护', style='List Bullet')

    doc.add_heading('技术难点', level=2)
    doc.add_paragraph('难点 1：如何确保adc采样进度能支撑mppt调节', style='List Bullet')
    doc.add_paragraph('难点 2：mppt开环做最外环加上闭环电压电流环如何调节出稳定的波形', style='List Bullet')

    # --- Section 4 ---
    doc.add_heading('四、解决方案与创新点', level=1)
    
    doc.add_heading('解决方案', level=2)
    doc.add_paragraph('针对问题 1 的解决方案：', style='List Bullet')
    doc.add_paragraph('针对问题 2 的解决方案：', style='List Bullet')

    doc.add_heading('创新点', level=2)
    doc.add_paragraph('创新点 1：', style='List Bullet')
    doc.add_paragraph('创新点 2：', style='List Bullet')
    doc.add_paragraph('创新点 3：', style='List Bullet')

    # --- Section 5 ---
    doc.add_heading('五、成果展示', level=1)
    
    doc.add_heading('已实现功能', level=2)
    doc.add_paragraph('功能 1：', style='List Bullet')
    doc.add_paragraph('功能 2：', style='List Bullet')
    doc.add_paragraph('功能 3：', style='List Bullet')

    doc.add_heading('性能 / 指标结果', level=2)
    doc.add_paragraph('指标 1：对比说明', style='List Bullet')
    doc.add_paragraph('指标 2：对比说明', style='List Bullet')

    # --- Section 6 ---
    doc.add_heading('六、效果与价值', level=1)
    
    doc.add_heading('项目效果', level=2)
    doc.add_paragraph('对业务的提升：', style='List Bullet')
    doc.add_paragraph('对系统稳定性的提升：', style='List Bullet')

    doc.add_heading('项目价值', level=2)
    doc.add_paragraph('技术价值：', style='List Bullet')
    doc.add_paragraph('业务价值：', style='List Bullet')
    doc.add_paragraph('推广价值：', style='List Bullet')

    doc.save(path)
    print(f"Updated docx created at: {path}")

if __name__ == "__main__":
    import os
    os.makedirs("input", exist_ok=True)
    create_dummy_docx("input/article.docx")
