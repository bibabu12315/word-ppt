# -*- coding: utf-8 -*-
"""
文件名称：app.py
主要作用：Streamlit Web 应用程序
实现功能：
1. 提供图形化界面供用户上传 Word 文档
2. 在线预览和编辑生成的 Markdown
3. 配置 PPT 生成参数
4. 下载生成的 PPT 文件
"""
import streamlit as st
import os
import json
import time
from docx import Document
from parser.markdown_parser import MarkdownParser
from parser.word_parser import WordParser
from parser.json_to_md import JsonToMdConverter
from ppt.generator import PPTGenerator
from utils.create_template import create_demo_template
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# 设置页面配置
st.set_page_config(
    page_title="Word转PPT 助手",
    page_icon="📊",
    layout="wide"
)

def main():
    # --- 标题居中 (紧凑版) ---
    st.markdown("""
        <style>
            .compact-title {
                text-align: center;
                margin-bottom: 0px !important;
                padding-bottom: 0px !important;
                padding-top: 0px !important;
                line-height: 1.2 !important;
            }
            .compact-subtitle {
                text-align: center;
                margin-top: 0px !important;
                margin-bottom: 5px !important;
                font-size: 0.9rem !important;
                color: #666;
            }
            /* 调整 Streamlit 默认的顶部 padding */
            .block-container {
                padding-top: 5rem !important;
                padding-bottom: 1rem !important;
            }
        </style>
        <h1 class="compact-title">📄 Word 转 PPT 助手</h1>
        <p class="compact-subtitle">上传 Word 文档，自动提取内容并生成 PPT。</p>
        <hr style="margin-top: 5px; margin-bottom: 10px;">
    """, unsafe_allow_html=True)

    # --- 侧边栏设置 ---
    with st.sidebar:
        st.header("设置")
        use_llm = st.toggle("使用 AI 增强 (LLM)", value=True, help="开启后将使用大模型优化内容，需要配置 API Key")
        
        st.divider()
        st.info("如果是首次运行，请确保已配置 .env 文件中的 API Key。")

    # 初始化 session state
    if "markdown_content" not in st.session_state:
        st.session_state.markdown_content = ""
    if "word_text" not in st.session_state:
        st.session_state.word_text = ""
    if "last_uploaded_file" not in st.session_state:
        st.session_state.last_uploaded_file = None

    # 准备路径
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_dir = os.path.join(base_dir, "input")
    output_dir = os.path.join(base_dir, "output")
    build_dir = os.path.join(base_dir, "build")
    
    # 确保目录存在
    os.makedirs(input_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(build_dir, exist_ok=True)

    # 定义文件路径
    input_docx_path = os.path.join(input_dir, "uploaded_article.docx")
    output_json_path = os.path.join(build_dir, "article.json")
    generated_md_path = os.path.join(input_dir, "generated.md")
    template_pptx_path = os.path.join(input_dir, "template.pptx")
    output_pptx_path = os.path.join(output_dir, "result.pptx")

    # 检查模板
    if not os.path.exists(template_pptx_path):
        create_demo_template(template_pptx_path)

    # --- 状态判断：是否已上传文件 ---
    has_file = st.session_state.last_uploaded_file is not None

    if not has_file:
        # --- 初始界面：居中上传 ---
        col_spacer1, col_center, col_spacer2 = st.columns([1, 2, 1])
        with col_center:
            st.info("👋 欢迎使用！请先上传一个 Word 文档开始。")
            uploaded_file = st.file_uploader("请上传 Word 文档 (.docx)", type=["docx"], key="uploader_center")
            
            if uploaded_file is not None:
                # 不再保存到本地，直接使用内存中的文件对象
                
                st.toast(f"已上传: {uploaded_file.name}")
                with st.spinner("正在解析文档并生成 Markdown..."):
                    try:
                        # 1. 提取 Word 纯文本
                        # 注意：Document 读取后指针会移动，下次读取前需要 seek(0)
                        doc = Document(uploaded_file)
                        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        st.session_state.word_text = full_text

                        # 2. Word -> JSON
                        uploaded_file.seek(0) # 重置文件指针
                        word_parser = WordParser()
                        data = word_parser.parse(uploaded_file)
                        word_parser.save_json(data, output_json_path)

                        # 3. JSON -> Markdown
                        converter = JsonToMdConverter()
                        with open(output_json_path, 'r', encoding='utf-8') as f:
                            json_data = json.load(f)
                        
                        mode = 1 if use_llm else 0
                        md_content = converter.convert(json_data, mode=mode)
                        
                        # 更新内容
                        st.session_state.markdown_content = md_content
                        st.session_state.last_uploaded_file = uploaded_file.name
                        
                        st.rerun()
                    except Exception as e:
                        st.error(f"解析失败: {str(e)}")
                        st.stop()

    else:
        # --- 已上传界面：左右分栏 ---
        col1, col2 = st.columns(2)
        
        # 左侧列
        with col1:
            st.subheader("📄 Word 原文预览")
            st.text_area(
                "Word Content", 
                value=st.session_state.word_text, 
                height=800, 
                disabled=True, 
                label_visibility="collapsed"
            )
            st.markdown("---")
            # 左下角的上传组件 (用于更换文件)
            uploaded_file_side = st.file_uploader("更换 Word 文档", type=["docx"], key="uploader_side")

        # 处理更换文件逻辑
        if uploaded_file_side is not None:
            if uploaded_file_side.name != st.session_state.last_uploaded_file:
                # 不再保存到本地
                
                st.toast(f"已上传新文件: {uploaded_file_side.name}")

                with st.spinner("正在重新解析..."):
                    try:
                        doc = Document(uploaded_file_side)
                        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        st.session_state.word_text = full_text

                        uploaded_file_side.seek(0) # 重置指针
                        word_parser = WordParser()
                        data = word_parser.parse(uploaded_file_side)
                        word_parser.save_json(data, output_json_path)

                        converter = JsonToMdConverter()
                        with open(output_json_path, 'r', encoding='utf-8') as f:
                            json_data = json.load(f)
                        
                        mode = 1 if use_llm else 0
                        md_content = converter.convert(json_data, mode=mode)
                        
                        st.session_state.markdown_content = md_content
                        st.session_state.last_uploaded_file = uploaded_file_side.name
                        
                        # 更新版本号
                        if "editor_version" not in st.session_state:
                            st.session_state.editor_version = 0
                        st.session_state.editor_version += 1
                        
                        st.rerun()
                    except Exception as e:
                        st.error(f"解析失败: {str(e)}")
                        st.stop()

        # 右侧列
        with col2:
            # 使用 Tabs 切换编辑和预览
            tab_editor, tab_preview = st.tabs(["📝 Markdown 编辑", "👀 幻灯片预览"])
            
            with tab_editor:
                if "editor_version" not in st.session_state:
                    st.session_state.editor_version = 0
                
                editor_key = f"editor_{st.session_state.last_uploaded_file}_{st.session_state.editor_version}"
                
                new_content = st.text_area(
                    "Markdown Editor", 
                    value=st.session_state.markdown_content,
                    height=800,
                    label_visibility="collapsed",
                    key=editor_key
                )
                
                # 实时同步用户的修改回 session_state
                if new_content != st.session_state.markdown_content:
                    st.session_state.markdown_content = new_content

            with tab_preview:
                if st.session_state.markdown_content:
                    try:
                        # 实时解析 Markdown 用于预览
                        md_parser = MarkdownParser()
                        # 将文本按行分割
                        lines = st.session_state.markdown_content.split('\n')
                        ppt_data = md_parser.parse_lines(lines)
                        
                        # --- 渲染预览 ---
                        st.markdown(f"### 封面: {ppt_data.cover_title}")
                        
                        # 渲染元数据 (meta_info)
                        if ppt_data.meta_info:
                            st.markdown("**封面信息:**")
                            for k, v in ppt_data.meta_info.items():
                                st.text(f"{k}: {v}")
                        
                        st.divider()
                        
                        for i, slide in enumerate(ppt_data.slides):
                            # 模拟 PPT 幻灯片的样式
                            with st.container(border=True):
                                st.markdown(f"#### 第 {i+1} 页: {slide.title}")
                                if slide.description:
                                    st.caption(slide.description)
                                
                                # 渲染内容块
                                for block in slide.blocks:
                                    cols = st.columns([1, 3])
                                    with cols[0]:
                                        if block.subtitle:
                                            st.markdown(f"**{block.subtitle}**")
                                    with cols[1]:
                                        for bullet in block.bullets:
                                            st.markdown(f"- {bullet}")
                                    
                                    if block.keyword:
                                        st.caption(f"🔑 关键词: {block.keyword}")
                                        
                    except Exception as e:
                        st.error(f"预览生成失败: {str(e)}")
                else:
                    st.info("暂无内容，请上传文档或在编辑区输入。")

            st.markdown("---")
            generate_btn = st.button("🚀 生成 PPT", type="primary", use_container_width=True)

        # 2. 处理生成 PPT
        if generate_btn:
            if not st.session_state.markdown_content:
                st.warning("请先上传文档或输入 Markdown 内容")
            else:
                progress_bar = st.progress(0)
                status_text = st.empty()

                try:
                    # 保存当前编辑框中的 Markdown 内容
                    with open(generated_md_path, 'w', encoding='utf-8') as f:
                        f.write(st.session_state.markdown_content)
                    
                    progress_bar.progress(30)
                    status_text.text("正在生成 PPT 文件...")

                    # Markdown -> PPT
                    md_parser = MarkdownParser()
                    presentation_data = md_parser.parse_file(generated_md_path)
                    
                    # 强制限制章节数量为 8
                    if len(presentation_data.slides) > 8:
                        st.warning(f"⚠️ 生成的章节数量 ({len(presentation_data.slides)}) 超过限制，已自动截取前 8 章。")
                        presentation_data.slides = presentation_data.slides[:8]
                    
                    generator = PPTGenerator(template_pptx_path, output_pptx_path)
                    generator.generate(presentation_data)
                    
                    progress_bar.progress(100)
                    status_text.text("✅ 转换完成！")
                    
                    # --- 下载按钮 (生成成功后显示) ---
                    with open(output_pptx_path, "rb") as f:
                        st.download_button(
                            label="📥 下载生成的 PPT",
                            data=f,
                            file_name="generated_presentation.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            use_container_width=True
                        )

                except Exception as e:
                    st.error(f"发生错误: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())

if __name__ == "__main__":
    main()
